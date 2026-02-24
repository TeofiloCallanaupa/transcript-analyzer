import docx
import csv
import sys
import os
import re

def process_docx_files(input_source, output_csv_path, log_callback=print, speaker_list=None, progress_callback=None, file_callback=None):
    """
    Processes all .docx files in a given folder OR a list of specific files, extracts interview transcripts,
    and combines them into a single .csv file with 'name', 'timestamp', 'statement' columns.

    input_source: Can be a string (folder path) or a list of strings (file paths).
    speaker_list: Optional list of speaker names to identify. If None, defaults to generic or empty.
    progress_callback: Optional callable(processed_paragraphs, total_paragraphs) for progress reporting.
    file_callback: Optional callable(file_index, total_files, filename) called when each file starts.
    """
    
    if speaker_list:
        SPEAKER_NAMES = speaker_list
    else:
        # Fallback if no list provided (though UI should always provide one)
        # Using a minimal default or keeping the old hardcoded list as a backup is an option,
        # but better to rely on config.
        SPEAKER_NAMES = ["Interviewer", "Respondent"] 

    # Create a regex pattern that matches any of the speaker names at the beginning of a line,
    # followed by an optional timestamp and then the statement.
    # To avoid ambiguity with statements starting with a name, we enforce that:
    # 1. The name is followed by a timestamp (and optional statement).
    # 2. OR The name is followed by the end of the line (indicating a missing timestamp).
    speaker_pattern = re.compile(
        r"^\s*(" + "|".join(re.escape(name) for name in SPEAKER_NAMES) + r")(?:(?:\s+((\d{1,2}:)?\d{2}:\d{2})\s*(.*))|(?:\s*$))"
    )

    docx_files = []
    if isinstance(input_source, str):
        if not os.path.isdir(input_source):
             log_callback(f"Error: The folder '{input_source}' does not exist or is not a directory.")
             return
        for filename in os.listdir(input_source):
            if filename.endswith(".docx"):
                docx_files.append(os.path.join(input_source, filename))
    elif isinstance(input_source, list):
        docx_files = [f for f in input_source if f.endswith(".docx")]
    
    if not docx_files:
        log_callback(f"No .docx files found in '{input_source}'. No CSV will be created.")
        return

    processed_transcript_data = []

    # Count total lines upfront for progress reporting.
    # Lines, not paragraphs — some DOCX files use soft breaks (shift+enter)
    # which put all content in a single paragraph.
    total_lines = 0
    if progress_callback:
        for docx_path in docx_files:
            try:
                doc = docx.Document(docx_path)
                for p in doc.paragraphs:
                    total_lines += len(p.text.split('\n'))
            except Exception:
                pass  # Will be handled in the main loop
    
    lines_done = 0

    for file_idx, docx_path in enumerate(docx_files):
        if file_callback:
            file_callback(file_idx + 1, len(docx_files), os.path.basename(docx_path))
        log_callback(f"Processing '{docx_path}'...")
        try:
            document = docx.Document(docx_path)
            current_speaker = None
            current_timestamp = None
            
            # Use os.path.basename for the filename column to keep it clean
            filename_base = os.path.basename(docx_path)

            for paragraph in document.paragraphs:
                lines = paragraph.text.split('\n')
                for line in lines:
                    text = line.strip()
                    
                    lines_done += 1
                    if progress_callback and total_lines > 0:
                        progress_callback(lines_done, total_lines)
                    
                    if not text:
                        continue

                    match = speaker_pattern.match(text)
                    if match:
                        # A new speaker line is found
                        current_speaker = match.group(1).strip()
                        current_timestamp = match.group(2).strip() if match.group(2) else ""
                        statement_part = match.group(4).strip() if match.group(4) else ""
                        if statement_part: # Only add if there is actual statement content
                            processed_transcript_data.append([filename_base, current_speaker, current_timestamp, statement_part])
                    else:
                        # This line is a continuation of the current speaker's statement
                        if current_speaker is not None:
                            processed_transcript_data.append([filename_base, current_speaker, current_timestamp, text])
                        # else: This line is not associated with any speaker, ignore it (e.g., initial garbage text)

        except Exception as e:
            log_callback(f"Error opening or reading .docx file '{docx_path}': {e}")
            # Continue processing other files even if one fails

    if not processed_transcript_data:
        speaker_list_str = ", ".join(SPEAKER_NAMES[:5])
        if len(SPEAKER_NAMES) > 5:
            speaker_list_str += f", ... ({len(SPEAKER_NAMES)} total)"
        log_callback(
            f"⚠️ No statements were found in your transcript(s). "
            f"This usually means the speaker names in your Settings don't match "
            f"the names in your documents."
        )
        log_callback(
            f"   Your current speaker names: {speaker_list_str}"
        )
        log_callback(
            f"   💡 To fix: Click the ⚙️ Settings icon and update the speaker names "
            f"to match exactly how they appear in your DOCX files."
        )
        return

    try:
        with open(output_csv_path, 'w', newline='', encoding='utf-8') as csvfile:
            csv_writer = csv.writer(csvfile)
            csv_writer.writerow(["source_file", "name", "timestamp", "statement"]) # Write header
            csv_writer.writerows(processed_transcript_data)
        log_callback(f"Successfully combined text from {len(docx_files)} files to '{output_csv_path}'.")
    except Exception as e:
        log_callback(f"Error writing to .csv file: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python docx_to_csv.py <input_folder_path> <output_csv_file>")
        sys.exit(1)

    input_folder = sys.argv[1]
    output_csv = sys.argv[2]
    process_docx_files(input_folder, output_csv)
