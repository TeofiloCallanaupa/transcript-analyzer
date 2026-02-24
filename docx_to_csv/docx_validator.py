"""
DOCX File Validator

Validates DOCX files before conversion to catch problems early 
and give users clear, actionable error messages.
"""

import os
import re
from dataclasses import dataclass, field

try:
    import docx
except ImportError:
    docx = None


@dataclass
class ValidationResult:
    """Result of validating a single DOCX file."""
    is_valid: bool
    file_path: str
    errors: list = field(default_factory=list)    # Blocking issues
    warnings: list = field(default_factory=list)  # Non-blocking but noteworthy


# Characters that could corrupt CSV output or indicate file problems.
# Includes ASCII control chars (except tab, newline, carriage return)
# and some problematic Unicode characters.
_CONTROL_CHAR_PATTERN = re.compile(
    r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]'
)


def validate_docx_file(file_path: str, speaker_list: list[str] | None = None) -> ValidationResult:
    """
    Validate a single DOCX file for structural and content issues.

    Runs checks in order, short-circuiting on structural failures:
    1. File exists and is readable
    2. File is a valid DOCX (ZIP-based Office Open XML)
    3. File is not empty (has at least one paragraph with text)
    4. File contains at least one line matching a configured speaker name
    5. Check for unusual/control characters that could break CSV output

    Args:
        file_path: Absolute path to the DOCX file.
        speaker_list: List of expected speaker names. If None/empty, check 4 is skipped.

    Returns:
        ValidationResult with is_valid, errors, and warnings.
    """
    result = ValidationResult(is_valid=True, file_path=file_path)

    # --- Check 1: File exists and is readable ---
    if not os.path.exists(file_path):
        result.is_valid = False
        result.errors.append("File not found or can't be read.")
        return result

    if not os.path.isfile(file_path):
        result.is_valid = False
        result.errors.append("This path is not a file.")
        return result

    if not os.access(file_path, os.R_OK):
        result.is_valid = False
        result.errors.append("File can't be read — check file permissions.")
        return result

    # --- Check 2: File is a valid DOCX ---
    if docx is None:
        result.is_valid = False
        result.errors.append("python-docx library is not installed.")
        return result

    try:
        document = docx.Document(file_path)
    except Exception:
        result.is_valid = False
        result.errors.append(
            "This file is not a valid DOCX document. "
            "It may be corrupted or renamed from another format (e.g. a .doc or .txt file)."
        )
        return result

    # --- Check 3: File is not empty ---
    all_text_lines = []
    for paragraph in document.paragraphs:
        lines = paragraph.text.split('\n')
        for line in lines:
            stripped = line.strip()
            if stripped:
                all_text_lines.append(stripped)

    # Also check table cells — some transcripts use table layouts
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    lines = paragraph.text.split('\n')
                    for line in lines:
                        stripped = line.strip()
                        if stripped:
                            all_text_lines.append(stripped)

    if not all_text_lines:
        result.is_valid = False
        result.errors.append("This document is empty — no text was found.")
        return result

    # --- Check 4: Speaker name matching ---
    if speaker_list:
        speaker_pattern = re.compile(
            r"^\s*(" + "|".join(re.escape(name) for name in speaker_list) + r")(?:(?:\s+((\d{1,2}:)?\d{2}:\d{2})\s*(.*))|(?:\s*$))"
        )
        has_speaker_match = any(speaker_pattern.match(line) for line in all_text_lines)

        if not has_speaker_match:
            speaker_preview = ", ".join(speaker_list[:5])
            if len(speaker_list) > 5:
                speaker_preview += f", ... ({len(speaker_list)} total)"

            # Show a sample of the first few lines so the user can see the mismatch
            sample_lines = all_text_lines[:3]
            sample_preview = "; ".join(f'"{line[:50]}"' for line in sample_lines)

            result.warnings.append(
                f"No speaker names were found in this document. "
                f"Your configured speakers are: {speaker_preview}. "
                f"First lines of the document: {sample_preview}. "
                f"Check that the names in Settings (⚙️) match your document exactly."
            )

    # --- Check 5: Control / unusual characters ---
    lines_with_issues = []
    for i, line in enumerate(all_text_lines):
        match = _CONTROL_CHAR_PATTERN.search(line)
        if match:
            char_hex = f"0x{ord(match.group()):02x}"
            preview = line[:60] + ("..." if len(line) > 60 else "")
            lines_with_issues.append((i + 1, char_hex, preview))

    if lines_with_issues:
        # Report up to 3 problem lines
        for line_num, char_hex, preview in lines_with_issues[:3]:
            result.warnings.append(
                f"Line {line_num} contains an unusual control character ({char_hex}): \"{preview}\""
            )
        if len(lines_with_issues) > 3:
            result.warnings.append(
                f"... and {len(lines_with_issues) - 3} more line(s) with unusual characters."
            )

    return result


def validate_docx_files(file_paths: list[str], speaker_list: list[str] | None = None) -> list[ValidationResult]:
    """
    Validate multiple DOCX files.

    Args:
        file_paths: List of absolute paths to DOCX files.
        speaker_list: List of expected speaker names.

    Returns:
        List of ValidationResult, one per file (same order as input).
    """
    return [validate_docx_file(path, speaker_list) for path in file_paths]
