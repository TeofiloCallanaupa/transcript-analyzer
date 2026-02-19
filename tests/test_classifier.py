import unittest
import os
import csv
import json
import tempfile
import shutil
from unittest.mock import patch, MagicMock
from csv_classifier import generate_prompt, classify_text_with_llm, process_csv_with_llm


# ============================================================================
# GENERATE_PROMPT TESTS
# ============================================================================

class TestGeneratePromptStructure(unittest.TestCase):
    """Test basic prompt structure with all arguments provided."""

    def test_contains_custom_categories(self):
        prompt = generate_prompt(["Happy", "Sad"], "How are you?", "I feel great.")
        self.assertIn("Happy", prompt)
        self.assertIn("Sad", prompt)

    def test_contains_context_and_text(self):
        prompt = generate_prompt(["Cat1"], "The question", "The answer")
        self.assertIn("The question", prompt)
        self.assertIn("The answer", prompt)

    def test_contains_custom_system_instruction(self):
        prompt = generate_prompt(["Cat1"], "ctx", "txt",
                                  system_instruction="You are a helpful assistant.")
        self.assertIn("You are a helpful assistant.", prompt)

    def test_contains_output_format_instructions(self):
        prompt = generate_prompt(["Cat1"], "ctx", "txt")
        self.assertIn("subthemes", prompt)
        self.assertIn("rationale", prompt)
        self.assertIn("JSON", prompt)

    def test_prompt_returns_string(self):
        result = generate_prompt(["Cat1"], "ctx", "txt")
        self.assertIsInstance(result, str)


class TestGeneratePromptDefaultBehavior(unittest.TestCase):
    """Test defaults and fallbacks in prompt generation."""

    def test_empty_categories_defaults_to_uncategorized(self):
        prompt = generate_prompt([], "ctx", "txt")
        self.assertIn("Uncategorized", prompt)

    def test_none_categories_defaults_to_uncategorized(self):
        prompt = generate_prompt(None, "ctx", "txt")
        self.assertIn("Uncategorized", prompt)

    def test_no_system_instruction_uses_default(self):
        prompt = generate_prompt(["Cat1"], "ctx", "txt", system_instruction=None)
        self.assertIn("expert text analyst", prompt)

    def test_empty_string_instruction_uses_default(self):
        prompt = generate_prompt(["Cat1"], "ctx", "txt", system_instruction="")
        self.assertIn("expert text analyst", prompt)

    def test_custom_instruction_replaces_default(self):
        prompt = generate_prompt(["Cat1"], "ctx", "txt",
                                  system_instruction="Custom instruction here.")
        self.assertIn("Custom instruction here.", prompt)
        self.assertNotIn("expert text analyst", prompt)


class TestGeneratePromptEdgeCases(unittest.TestCase):
    """Test edge cases in prompt generation."""

    def test_single_category(self):
        prompt = generate_prompt(["OnlyOne"], "ctx", "txt")
        self.assertIn("OnlyOne", prompt)

    def test_many_categories(self):
        cats = [f"Category_{i}" for i in range(20)]
        prompt = generate_prompt(cats, "ctx", "txt")
        for cat in cats:
            self.assertIn(cat, prompt)

    def test_category_with_special_characters(self):
        prompt = generate_prompt(["Culture & Engagement", "HR/Talent", "Risk (High)"], "ctx", "txt")
        self.assertIn("Culture & Engagement", prompt)
        self.assertIn("HR/Talent", prompt)
        self.assertIn("Risk (High)", prompt)

    def test_empty_context(self):
        prompt = generate_prompt(["Cat1"], "", "Some text")
        self.assertIn("Some text", prompt)

    def test_empty_text(self):
        prompt = generate_prompt(["Cat1"], "Some context", "")
        self.assertIn("Some context", prompt)

    def test_none_context(self):
        # Context could be None from certain code paths
        prompt = generate_prompt(["Cat1"], None, "txt")
        self.assertIn("None", prompt)  # It gets interpolated as the string "None"
        self.assertIsInstance(prompt, str)

    def test_multiline_text_input(self):
        text = "Line one.\nLine two.\nLine three."
        prompt = generate_prompt(["Cat1"], "ctx", text)
        self.assertIn("Line one.", prompt)
        self.assertIn("Line three.", prompt)

    def test_very_long_system_instruction(self):
        long_instruction = "Analyze carefully. " * 100
        prompt = generate_prompt(["Cat1"], "ctx", "txt",
                                  system_instruction=long_instruction)
        self.assertIn("Analyze carefully.", prompt)
        self.assertNotIn("expert text analyst", prompt)

    def test_uncategorized_fallback_instructions_present(self):
        """Prompt should explain what to do when no category matches."""
        prompt = generate_prompt(["Cat1"], "ctx", "txt")
        self.assertIn("Uncategorized", prompt)
        self.assertIn("does not match", prompt)

    def test_categories_formatted_as_bullet_points(self):
        """Each category should appear as a bullet point."""
        prompt = generate_prompt(["Alpha", "Beta", "Gamma"], "ctx", "txt")
        self.assertIn("• Alpha", prompt)
        self.assertIn("• Beta", prompt)
        self.assertIn("• Gamma", prompt)


# ============================================================================
# CLASSIFY_TEXT_WITH_LLM TESTS (with mocks — no real API calls)
# ============================================================================

class TestClassifyTextNoApiKey(unittest.TestCase):
    """Test that classify_text_with_llm raises an error when no API key is available."""

    @patch.dict(os.environ, {}, clear=True)
    @patch('csv_classifier.load_dotenv')
    def test_raises_value_error_without_api_key(self, mock_dotenv):
        """Should raise ValueError when no API key is provided or in env."""
        logs = []
        with self.assertRaises(ValueError) as ctx:
            classify_text_with_llm("some text", log_callback=logs.append)
        self.assertIn("OPENAI_API_KEY", str(ctx.exception))

    @patch.dict(os.environ, {}, clear=True)
    @patch('csv_classifier.load_dotenv')
    def test_logs_error_message_without_api_key(self, mock_dotenv):
        """Should log an error message before raising."""
        logs = []
        try:
            classify_text_with_llm("some text", log_callback=logs.append)
        except ValueError:
            pass
        self.assertTrue(any("OPENAI_API_KEY" in msg for msg in logs))


class TestClassifyTextWithMockedApi(unittest.TestCase):
    """Test classify_text_with_llm with a mocked OpenAI API."""

    def _mock_openai_response(self, content):
        """Create a mock OpenAI API response."""
        mock_response = MagicMock()
        mock_response.choices = [MagicMock()]
        mock_response.choices[0].message.content = content
        return mock_response

    @patch('csv_classifier.openai')
    @patch('csv_classifier.load_dotenv')
    def test_valid_json_response_parsed(self, mock_dotenv, mock_openai):
        """A valid JSON response should be parsed and returned."""
        mock_openai.api_key = None
        response_json = json.dumps({
            "subthemes": ["Happy", "Positive"],
            "rationale": "The text expresses joy."
        })
        mock_openai.chat.completions.create.return_value = self._mock_openai_response(response_json)

        result = classify_text_with_llm("I feel great!", api_key="fake-key",
                                         categories=["Happy", "Sad"])
        self.assertEqual(result["subthemes"], ["Happy", "Positive"])
        self.assertEqual(result["rationale"], "The text expresses joy.")

    @patch('csv_classifier.openai')
    @patch('csv_classifier.load_dotenv')
    def test_invalid_json_returns_error(self, mock_dotenv, mock_openai):
        """If the LLM returns invalid JSON, should return ERROR gracefully."""
        mock_openai.api_key = None
        mock_openai.chat.completions.create.return_value = self._mock_openai_response(
            "This is not JSON at all"
        )
        logs = []
        result = classify_text_with_llm("text", api_key="fake-key", log_callback=logs.append)
        self.assertEqual(result["subthemes"], ["ERROR"])
        self.assertIn("JSON decoding error", result["rationale"])

    @patch('csv_classifier.openai')
    @patch('csv_classifier.load_dotenv')
    def test_missing_subthemes_key_returns_error(self, mock_dotenv, mock_openai):
        """JSON without 'subthemes' key should raise and return ERROR."""
        mock_openai.api_key = None
        response_json = json.dumps({"only_rationale": "no subthemes here"})
        mock_openai.chat.completions.create.return_value = self._mock_openai_response(response_json)

        logs = []
        result = classify_text_with_llm("text", api_key="fake-key", log_callback=logs.append)
        self.assertEqual(result["subthemes"], ["ERROR"])

    @patch('csv_classifier.openai')
    @patch('csv_classifier.load_dotenv')
    def test_missing_rationale_key_returns_error(self, mock_dotenv, mock_openai):
        """JSON without 'rationale' key should raise and return ERROR."""
        mock_openai.api_key = None
        response_json = json.dumps({"subthemes": ["Cat1"]})
        mock_openai.chat.completions.create.return_value = self._mock_openai_response(response_json)

        logs = []
        result = classify_text_with_llm("text", api_key="fake-key", log_callback=logs.append)
        self.assertEqual(result["subthemes"], ["ERROR"])

    @patch('csv_classifier.openai')
    @patch('csv_classifier.load_dotenv')
    def test_api_exception_returns_error(self, mock_dotenv, mock_openai):
        """If the API call throws an exception, should return ERROR gracefully."""
        mock_openai.api_key = None
        mock_openai.chat.completions.create.side_effect = Exception("Network error")

        logs = []
        result = classify_text_with_llm("text", api_key="fake-key", log_callback=logs.append)
        self.assertEqual(result["subthemes"], ["ERROR"])
        self.assertIn("Classification error", result["rationale"])

    @patch('csv_classifier.openai')
    @patch('csv_classifier.load_dotenv')
    def test_empty_subthemes_list_still_valid(self, mock_dotenv, mock_openai):
        """LLM might return empty subthemes list — it should still parse."""
        mock_openai.api_key = None
        response_json = json.dumps({
            "subthemes": [],
            "rationale": "Nothing matched."
        })
        mock_openai.chat.completions.create.return_value = self._mock_openai_response(response_json)

        result = classify_text_with_llm("text", api_key="fake-key")
        self.assertEqual(result["subthemes"], [])

    @patch('csv_classifier.openai')
    @patch('csv_classifier.load_dotenv')
    def test_response_with_extra_whitespace(self, mock_dotenv, mock_openai):
        """LLM responses often have leading/trailing whitespace and newlines."""
        mock_openai.api_key = None
        response_json = '\n  ' + json.dumps({
            "subthemes": ["Cat1"],
            "rationale": "Matched."
        }) + '  \n'
        mock_openai.chat.completions.create.return_value = self._mock_openai_response(response_json)

        result = classify_text_with_llm("text", api_key="fake-key")
        self.assertEqual(result["subthemes"], ["Cat1"])


# ============================================================================
# PROCESS_CSV_WITH_LLM TESTS (with mocks)
# ============================================================================

class TestProcessCsvWithLlm(unittest.TestCase):
    """Test the full CSV processing pipeline with mocked LLM calls."""

    def setUp(self):
        self.test_dir = tempfile.mkdtemp()
        self.csv_path = os.path.join(self.test_dir, "input.csv")
        self.logs = []

    def tearDown(self):
        shutil.rmtree(self.test_dir)

    def _log(self, msg):
        self.logs.append(msg)

    def _write_csv(self, rows):
        """Write a CSV file with header + data rows."""
        with open(self.csv_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["source_file", "name", "timestamp", "statement"])
            writer.writerows(rows)

    def test_nonexistent_csv_logs_error(self):
        """Processing a non-existent CSV should log an error and not crash."""
        process_csv_with_llm("/no/such/file.csv", log_callback=self._log)
        self.assertTrue(any("does not exist" in msg for msg in self.logs))

    @patch('csv_classifier.classify_text_with_llm')
    def test_interviewer_row_not_classified(self, mock_classify):
        """Rows from interviewers should NOT call the LLM."""
        self._write_csv([
            ["file.docx", "InterviewerM", "10:00", "What do you think?"],
        ])
        process_csv_with_llm(self.csv_path, api_key="fake", 
                            categories=["Cat1", "Interviewer"],
                            log_callback=self._log)
        mock_classify.assert_not_called()

    @patch('csv_classifier.classify_text_with_llm')
    def test_non_interviewer_row_classified(self, mock_classify):
        """Rows from non-interviewers SHOULD call the LLM."""
        mock_classify.return_value = {
            "subthemes": ["Cat1"],
            "rationale": "Matched."
        }
        self._write_csv([
            ["file.docx", "Laura", "10:01", "I think AI is great."],
        ])
        process_csv_with_llm(self.csv_path, api_key="fake",
                            categories=["Cat1"],
                            log_callback=self._log)
        mock_classify.assert_called_once()

    @patch('csv_classifier.classify_text_with_llm')
    def test_interviewer_context_passed_to_classifier(self, mock_classify):
        """The most recent interviewer statement should be passed as context."""
        mock_classify.return_value = {
            "subthemes": ["Cat1"],
            "rationale": "Matched."
        }
        self._write_csv([
            ["file.docx", "InterviewerM", "10:00", "Tell me about AI governance."],
            ["file.docx", "Laura", "10:01", "We need better policies."],
        ])
        process_csv_with_llm(self.csv_path, api_key="fake",
                            categories=["Cat1"],
                            log_callback=self._log)
        
        # The context passed should be the interviewer's question
        call_args = mock_classify.call_args
        self.assertEqual(call_args[0][1], "Tell me about AI governance.")

    @patch('csv_classifier.classify_text_with_llm')
    def test_context_resets_on_new_source_file(self, mock_classify):
        """Context should reset when processing rows from a different source file."""
        mock_classify.return_value = {
            "subthemes": ["Cat1"],
            "rationale": "Matched."
        }
        self._write_csv([
            ["file1.docx", "InterviewerM", "10:00", "Question from file 1"],
            ["file2.docx", "Laura", "10:01", "Answer in file 2"],
        ])
        process_csv_with_llm(self.csv_path, api_key="fake",
                            categories=["Cat1"],
                            log_callback=self._log)
        
        # Context should have reset, not carry over from file1
        call_args = mock_classify.call_args
        context = call_args[0][1]
        self.assertIn("No preceding question", context)

    @patch('csv_classifier.classify_text_with_llm')
    def test_output_csv_has_category_columns(self, mock_classify):
        """The output CSV should have category columns + Rationale appended."""
        mock_classify.return_value = {
            "subthemes": ["Governance"],
            "rationale": "Discusses policy."
        }
        self._write_csv([
            ["file.docx", "Laura", "10:01", "We need governance."],
        ])
        categories = ["Governance", "Trust", "Interviewer"]
        process_csv_with_llm(self.csv_path, api_key="fake",
                            categories=categories,
                            log_callback=self._log)
        
        # Read the updated CSV
        with open(self.csv_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            header = next(reader)
            data_row = next(reader)
        
        # Header should have original 4 + categories + Rationale
        self.assertEqual(len(header), 4 + len(categories) + 1)
        self.assertIn("Governance", header)
        self.assertIn("Trust", header)
        self.assertIn("Rationale", header)

    @patch('csv_classifier.classify_text_with_llm')
    def test_category_scores_are_binary(self, mock_classify):
        """Category columns should be 1 (matched) or 0 (not matched)."""
        mock_classify.return_value = {
            "subthemes": ["Cat_A"],
            "rationale": "Matched Cat_A."
        }
        self._write_csv([
            ["file.docx", "Laura", "10:01", "Some statement"],
        ])
        categories = ["Cat_A", "Cat_B", "Cat_C"]
        process_csv_with_llm(self.csv_path, api_key="fake",
                            categories=categories,
                            log_callback=self._log)
        
        with open(self.csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            row = next(reader)
        
        self.assertEqual(row['Cat_A'], '1')
        self.assertEqual(row['Cat_B'], '0')
        self.assertEqual(row['Cat_C'], '0')

    @patch('csv_classifier.classify_text_with_llm')
    def test_short_rows_passed_through(self, mock_classify):
        """Rows with fewer than 4 columns should be written as-is without classification."""
        # Manually write a CSV with a short row
        with open(self.csv_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["source_file", "name", "timestamp", "statement"])
            writer.writerow(["only_two_cols", "name"])  # Short row
            writer.writerow(["file.docx", "Laura", "10:01", "Real statement"])
        
        mock_classify.return_value = {
            "subthemes": ["Cat1"],
            "rationale": "Matched."
        }
        process_csv_with_llm(self.csv_path, api_key="fake",
                            categories=["Cat1"],
                            log_callback=self._log)
        
        # Should still process without crashing
        with open(self.csv_path, 'r', encoding='utf-8') as f:
            rows = list(csv.reader(f))
        self.assertTrue(len(rows) >= 2)  # header + at least one data row

    @patch('csv_classifier.classify_text_with_llm')
    def test_empty_categories_defaults(self, mock_classify):
        """When no categories are provided, should default to ['Uncategorized']."""
        mock_classify.return_value = {
            "subthemes": ["Uncategorized"],
            "rationale": "No match."
        }
        self._write_csv([
            ["file.docx", "Laura", "10:01", "Some text"],
        ])
        process_csv_with_llm(self.csv_path, api_key="fake",
                            categories=None,
                            log_callback=self._log)
        
        with open(self.csv_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            header = next(reader)
        self.assertIn("Uncategorized", header)

    @patch('csv_classifier.classify_text_with_llm')
    def test_interviewer_category_marked_for_interviewer_rows(self, mock_classify):
        """When 'Interviewer' is in categories, interviewer rows should get score 1."""
        self._write_csv([
            ["file.docx", "InterviewerM", "10:00", "A question"],
        ])
        categories = ["Cat1", "Interviewer"]
        process_csv_with_llm(self.csv_path, api_key="fake",
                            categories=categories,
                            log_callback=self._log)
        
        with open(self.csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            row = next(reader)
        
        self.assertEqual(row['Interviewer'], '1')
        self.assertEqual(row['Cat1'], '0')
        self.assertEqual(row['Rationale'], 'Statement from interviewer.')


# ============================================================================
# INTEGRATION: DOCX → CSV → CLASSIFIER PIPELINE
# ============================================================================

class TestEndToEndPipeline(unittest.TestCase):
    """Test the complete pipeline: DOCX → CSV → Classification (with mock LLM)."""

    def setUp(self):
        self.test_dir = tempfile.mkdtemp()
        self.logs = []
    
    def tearDown(self):
        shutil.rmtree(self.test_dir)

    def _log(self, msg):
        self.logs.append(msg)

    @patch('csv_classifier.classify_text_with_llm')
    def test_docx_to_csv_to_classification(self, mock_classify):
        """Full pipeline: create DOCX → convert to CSV → classify with mocked LLM."""
        from docx import Document
        from docx_to_csv.docx_to_csv import process_docx_files

        mock_classify.return_value = {
            "subthemes": ["Governance"],
            "rationale": "Discusses governance."
        }

        # Step 1: Create a DOCX
        docx_path = os.path.join(self.test_dir, "interview.docx")
        doc = Document()
        doc.add_paragraph("InterviewerM 10:00 What about governance?")
        doc.add_paragraph("Laura 10:01 We need better AI governance frameworks.")
        doc.save(docx_path)

        # Step 2: Convert to CSV
        csv_path = os.path.join(self.test_dir, "output.csv")
        process_docx_files([docx_path], csv_path, log_callback=self._log,
                           speaker_list=["InterviewerM", "Laura"])
        
        self.assertTrue(os.path.exists(csv_path))
        with open(csv_path, 'r', encoding='utf-8') as f:
            rows = list(csv.DictReader(f))
        self.assertEqual(len(rows), 2)

        # Step 3: Classify
        categories = ["Governance", "Trust", "Interviewer"]
        process_csv_with_llm(csv_path, api_key="fake",
                            categories=categories,
                            log_callback=self._log)

        # Verify final output
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            final_rows = list(reader)
        
        self.assertEqual(len(final_rows), 2)
        # Interviewer row should be marked
        self.assertEqual(final_rows[0]['Interviewer'], '1')
        # Laura's row should have Governance marked
        self.assertEqual(final_rows[1]['Governance'], '1')
        self.assertEqual(final_rows[1]['Trust'], '0')


if __name__ == '__main__':
    unittest.main()
