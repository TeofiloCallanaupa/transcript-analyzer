import unittest
import os
import shutil
import tempfile
import csv
from docx import Document
from docx_to_csv.docx_to_csv import process_docx_files


class DocxTestBase(unittest.TestCase):
    """Base class with helper methods for DOCX tests."""

    def setUp(self):
        self.test_dir = tempfile.mkdtemp()
        self.output_csv = os.path.join(self.test_dir, "output.csv")
        self.log_messages = []

    def tearDown(self):
        shutil.rmtree(self.test_dir)

    def _log(self, msg):
        self.log_messages.append(msg)

    def _create_docx(self, filename, paragraphs):
        """Create a DOCX file with the given paragraphs (list of strings).
        Each string becomes a separate paragraph in the document."""
        path = os.path.join(self.test_dir, filename)
        doc = Document()
        for para in paragraphs:
            doc.add_paragraph(para)
        doc.save(path)
        return path

    def _create_docx_soft_breaks(self, filename, lines):
        """Create a DOCX file where all lines are in a single paragraph (soft breaks/shift+enter)."""
        path = os.path.join(self.test_dir, filename)
        doc = Document()
        doc.add_paragraph('\n'.join(lines))
        doc.save(path)
        return path

    def _read_csv_rows(self):
        """Read the output CSV and return list of dicts."""
        with open(self.output_csv, 'r', encoding='utf-8') as f:
            return list(csv.DictReader(f))

    def _read_csv_raw(self):
        """Read the output CSV and return list of lists (including header)."""
        with open(self.output_csv, 'r', encoding='utf-8') as f:
            return list(csv.reader(f))


# ============================================================================
# BASIC EXTRACTION
# ============================================================================

class TestBasicExtraction(DocxTestBase):
    """Test basic speaker/timestamp/statement extraction."""

    def test_single_speaker_single_statement(self):
        path = self._create_docx("test.docx", [
            "Alice 10:00 Hello world"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]['name'], 'Alice')
        self.assertEqual(rows[0]['timestamp'], '10:00')
        self.assertEqual(rows[0]['statement'], 'Hello world')

    def test_multiple_speakers(self):
        path = self._create_docx("test.docx", [
            "Alice 10:00 First statement",
            "Bob 10:01 Second statement",
            "Alice 10:02 Third statement"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice", "Bob"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 3)
        self.assertEqual(rows[0]['name'], 'Alice')
        self.assertEqual(rows[1]['name'], 'Bob')
        self.assertEqual(rows[2]['name'], 'Alice')

    def test_source_file_column_uses_basename(self):
        path = self._create_docx("my_interview.docx", [
            "Alice 10:00 Testing"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(rows[0]['source_file'], 'my_interview.docx')

    def test_csv_header_columns(self):
        """Verify the CSV has the correct header columns."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 Testing headers"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        raw = self._read_csv_raw()
        self.assertEqual(raw[0], ["source_file", "name", "timestamp", "statement"])


# ============================================================================
# TIMESTAMP FORMATS
# ============================================================================

class TestTimestampFormats(DocxTestBase):
    """Test various timestamp formats the regex must handle."""

    def test_mm_ss_timestamp(self):
        path = self._create_docx("test.docx", ["Alice 05:30 Short timestamp"])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(rows[0]['timestamp'], '05:30')

    def test_h_mm_ss_timestamp(self):
        path = self._create_docx("test.docx", ["Alice 1:02:49 After an hour"])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(rows[0]['timestamp'], '1:02:49')

    def test_hh_mm_ss_timestamp(self):
        path = self._create_docx("test.docx", ["Alice 00:30:15 Leading zeros hour"])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(rows[0]['timestamp'], '00:30:15')

    def test_speaker_without_timestamp(self):
        """Speaker line with no timestamp (just the name alone on a line)."""
        path = self._create_docx_soft_breaks("test.docx", [
            "Alice",
            "This is a statement without a timestamp."
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]['name'], 'Alice')
        self.assertEqual(rows[0]['timestamp'], '')
        self.assertEqual(rows[0]['statement'], 'This is a statement without a timestamp.')

    def test_speaker_with_timestamp_but_no_statement_on_same_line(self):
        """Speaker line has a timestamp but no statement text on the same line.
        The statement comes on the next line as a continuation."""
        path = self._create_docx_soft_breaks("test.docx", [
            "Alice 10:00",
            "The actual statement is on the next line."
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        # The speaker line with just timestamp has no statement_part, so no row for it.
        # The next line is a continuation.
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]['name'], 'Alice')
        self.assertEqual(rows[0]['timestamp'], '10:00')
        self.assertEqual(rows[0]['statement'], 'The actual statement is on the next line.')


# ============================================================================
# CONTINUATION LINES & MULTI-LINE STATEMENTS
# ============================================================================

class TestContinuationLines(DocxTestBase):
    """Test multi-line statements that continue under a speaker."""

    def test_continuation_lines_with_soft_breaks(self):
        """When a speaker has multi-line text via soft breaks, each line becomes a row."""
        path = self._create_docx_soft_breaks("test.docx", [
            "Alice 10:00 First line of statement",
            "This is a continuation line.",
            "And another continuation.",
            "",
            "Bob 10:01 Bob's line"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice", "Bob"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 4)
        self.assertEqual(rows[0]['name'], 'Alice')
        self.assertEqual(rows[0]['statement'], 'First line of statement')
        self.assertEqual(rows[1]['name'], 'Alice')
        self.assertEqual(rows[1]['statement'], 'This is a continuation line.')
        self.assertEqual(rows[2]['name'], 'Alice')
        self.assertEqual(rows[2]['statement'], 'And another continuation.')
        self.assertEqual(rows[3]['name'], 'Bob')

    def test_continuation_preserves_speaker_timestamp(self):
        """Continuation lines should carry over the original speaker's timestamp."""
        path = self._create_docx_soft_breaks("test.docx", [
            "Alice 10:05 First part",
            "Second part continues",
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['timestamp'], '10:05')
        self.assertEqual(rows[1]['timestamp'], '10:05')  # Same timestamp carried over

    def test_continuation_across_separate_paragraphs(self):
        """Continuation lines work across separate paragraphs in the DOCX."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 First paragraph statement",
            "This continues Alice's thought in a new paragraph"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['name'], 'Alice')
        self.assertEqual(rows[1]['name'], 'Alice')
        self.assertEqual(rows[1]['statement'], "This continues Alice's thought in a new paragraph")


# ============================================================================
# MULTIPLE FILES
# ============================================================================

class TestMultipleFiles(DocxTestBase):
    """Test combining multiple DOCX files into one CSV."""

    def test_two_files_combined(self):
        path1 = self._create_docx("interview1.docx", [
            "Alice 10:00 From file one"
        ])
        path2 = self._create_docx("interview2.docx", [
            "Bob 11:00 From file two"
        ])
        process_docx_files([path1, path2], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice", "Bob"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['source_file'], 'interview1.docx')
        self.assertEqual(rows[1]['source_file'], 'interview2.docx')

    def test_folder_input(self):
        """Test passing a folder path (string) as input_source."""
        sub_dir = os.path.join(self.test_dir, "docs")
        os.makedirs(sub_dir)
        doc = Document()
        doc.add_paragraph("Alice 10:00 Folder test")
        doc.save(os.path.join(sub_dir, "test.docx"))

        process_docx_files(sub_dir, self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]['statement'], 'Folder test')

    def test_folder_ignores_non_docx_files(self):
        """When scanning a folder, only .docx files should be picked up."""
        sub_dir = os.path.join(self.test_dir, "mixed")
        os.makedirs(sub_dir)
        
        # Create a valid DOCX
        doc = Document()
        doc.add_paragraph("Alice 10:00 Valid content")
        doc.save(os.path.join(sub_dir, "valid.docx"))
        
        # Create non-docx files
        with open(os.path.join(sub_dir, "notes.txt"), 'w') as f:
            f.write("not a docx")
        with open(os.path.join(sub_dir, "data.csv"), 'w') as f:
            f.write("col1,col2")

        process_docx_files(sub_dir, self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]['statement'], 'Valid content')

    def test_three_files_all_combined_in_order(self):
        """Three files should all contribute rows to the CSV."""
        paths = []
        for i in range(3):
            p = self._create_docx(f"file{i}.docx", [
                f"Alice 10:0{i} Statement from file {i}"
            ])
            paths.append(p)
        process_docx_files(paths, self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 3)
        for i in range(3):
            self.assertEqual(rows[i]['source_file'], f'file{i}.docx')


# ============================================================================
# SPEAKER NAME EDGE CASES
# ============================================================================

class TestSpeakerNameEdgeCases(DocxTestBase):
    """Test tricky speaker name scenarios that non-technical users might encounter."""

    def test_speaker_name_with_special_regex_chars(self):
        """Speaker names with periods, parentheses etc. must be escaped in regex."""
        path = self._create_docx("test.docx", [
            "Dr. Smith 10:00 Hello from the doctor"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Dr. Smith"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]['name'], 'Dr. Smith')

    def test_speaker_name_substring_does_not_false_match(self):
        """'Al' should not match 'Alice'. Only exact names should match."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 I am Alice",
            "Al 10:01 I am Al",
        ])
        # Only "Alice" is in the speaker list, not "Al"
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        # "Al 10:01 I am Al" should be treated as a continuation of Alice's speech
        # because "Al" is not in the speaker list
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['name'], 'Alice')
        self.assertEqual(rows[0]['statement'], 'I am Alice')
        self.assertEqual(rows[1]['name'], 'Alice')  # continuation

    def test_both_similar_names_in_list(self):
        """When both 'Al' and 'Alice' are in the speaker list, both should match."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 I am Alice",
            "Al 10:01 I am Al",
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice", "Al"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['name'], 'Alice')
        self.assertEqual(rows[1]['name'], 'Al')

    def test_speaker_name_case_sensitivity(self):
        """Speaker names are case-sensitive. 'alice' ≠ 'Alice'."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 Correct case",
            "alice 10:01 Wrong case"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        # "alice" should NOT match, treated as continuation of Alice
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['name'], 'Alice')
        self.assertEqual(rows[1]['name'], 'Alice')  # "alice 10:01 Wrong case" is continuation

    def test_empty_speaker_list_uses_defaults(self):
        """An empty speaker list [] should fall back to defaults."""
        path = self._create_docx("test.docx", [
            "Interviewer 10:00 Default question",
            "Respondent 10:01 Default answer"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=[])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['name'], 'Interviewer')
        self.assertEqual(rows[1]['name'], 'Respondent')


# ============================================================================
# UNICODE AND SPECIAL CHARACTERS
# ============================================================================

class TestUnicodeAndSpecialChars(DocxTestBase):
    """Test that international characters and special content are handled correctly."""

    def test_unicode_in_statements(self):
        """Accented characters, non-Latin scripts should pass through."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 Héllo wörld café résumé"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(rows[0]['statement'], 'Héllo wörld café résumé')

    def test_emoji_in_statements(self):
        """Emojis should pass through without crashing."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 I love this project 🎉👍"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertIn("🎉", rows[0]['statement'])

    def test_commas_and_quotes_in_statements(self):
        """CSV-special characters (commas, quotes) in statements must be preserved."""
        path = self._create_docx("test.docx", [
            'Alice 10:00 She said, "hello, how are you?"'
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertIn("hello, how are you?", rows[0]['statement'])

    def test_unicode_speaker_names(self):
        """Speaker names with accented characters should work."""
        path = self._create_docx("test.docx", [
            "José 10:00 Hola mundo"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["José"])
        rows = self._read_csv_rows()
        self.assertEqual(rows[0]['name'], 'José')


# ============================================================================
# ERROR HANDLING AND EDGE CASES
# ============================================================================

class TestErrorHandling(DocxTestBase):
    """Test error handling — things non-technical users WILL run into."""

    def test_empty_docx_produces_no_csv(self):
        """An empty DOCX should not create an output CSV."""
        path = self._create_docx("empty.docx", [])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        self.assertFalse(os.path.exists(self.output_csv))

    def test_no_matching_speakers_produces_no_csv(self):
        """If no lines match any known speaker, no CSV is created."""
        path = self._create_docx("test.docx", [
            "UnknownPerson 10:00 This shouldn't match",
            "Some random text"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice", "Bob"])
        self.assertFalse(os.path.exists(self.output_csv))

    def test_text_before_first_speaker_is_ignored(self):
        """Lines before any known speaker should be ignored (preamble, headers, etc)."""
        path = self._create_docx_soft_breaks("test.docx", [
            "Interview Transcript - Company XYZ",
            "Date: January 15, 2026",
            "Confidential",
            "",
            "Alice 10:00 The real content starts here"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]['statement'], 'The real content starts here')

    def test_nonexistent_folder_logs_error(self):
        """Passing a non-existent folder path should log an error and not crash."""
        process_docx_files("/nonexistent/path", self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        self.assertFalse(os.path.exists(self.output_csv))
        self.assertTrue(any("does not exist" in msg for msg in self.log_messages))

    def test_non_docx_files_in_list_are_ignored(self):
        """Non-.docx files in the list should be filtered out."""
        txt_path = os.path.join(self.test_dir, "notes.txt")
        with open(txt_path, 'w') as f:
            f.write("not a docx")
        process_docx_files([txt_path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        self.assertFalse(os.path.exists(self.output_csv))

    def test_default_speaker_list_fallback(self):
        """When speaker_list is None, defaults to ['Interviewer', 'Respondent']."""
        path = self._create_docx("test.docx", [
            "Interviewer 10:00 Default question",
            "Respondent 10:01 Default answer"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=None)
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 2)
        self.assertEqual(rows[0]['name'], 'Interviewer')
        self.assertEqual(rows[1]['name'], 'Respondent')

    def test_one_bad_file_doesnt_stop_others(self):
        """If one DOCX is corrupt, the other files should still process."""
        good_path = self._create_docx("good.docx", [
            "Alice 10:00 Good file content"
        ])
        bad_path = os.path.join(self.test_dir, "bad.docx")
        with open(bad_path, 'w') as f:
            f.write("this is not a valid docx file")

        process_docx_files([bad_path, good_path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]['statement'], 'Good file content')
        self.assertTrue(any("Error" in msg for msg in self.log_messages))

    def test_all_files_corrupt_produces_no_csv(self):
        """If every DOCX is corrupt, no CSV should be created."""
        bad1 = os.path.join(self.test_dir, "bad1.docx")
        bad2 = os.path.join(self.test_dir, "bad2.docx")
        for path in [bad1, bad2]:
            with open(path, 'w') as f:
                f.write("not a docx")

        process_docx_files([bad1, bad2], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        self.assertFalse(os.path.exists(self.output_csv))

    def test_empty_list_input(self):
        """Passing an empty list should log and not crash."""
        process_docx_files([], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        self.assertFalse(os.path.exists(self.output_csv))

    def test_read_only_output_directory(self):
        """If the output path is not writable, should log error and not crash."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 Some content"
        ])
        bad_output = "/nonexistent_dir/output.csv"
        process_docx_files([path], bad_output, log_callback=self._log,
                           speaker_list=["Alice"])
        self.assertTrue(any("Error" in msg for msg in self.log_messages))

    def test_docx_with_only_whitespace_paragraphs(self):
        """A DOCX with only whitespace/blank paragraphs should produce no CSV."""
        path = self._create_docx("blank.docx", [
            "   ",
            "  ",
            "",
            "\t"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        self.assertFalse(os.path.exists(self.output_csv))

    def test_log_callback_receives_messages(self):
        """Verify log_callback is called with processing messages."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 Testing logs"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        self.assertTrue(len(self.log_messages) > 0)
        self.assertTrue(any("Processing" in msg for msg in self.log_messages))
        self.assertTrue(any("Successfully" in msg for msg in self.log_messages))


# ============================================================================
# REAL-WORLD SCENARIOS
# ============================================================================

class TestRealWorldScenarios(DocxTestBase):
    """Tests that simulate what non-technical users will actually encounter."""

    def test_full_interview_flow(self):
        """Simulate a full interview transcript with multiple speakers, timestamps, and continuations."""
        path = self._create_docx_soft_breaks("real_interview.docx", [
            "InterviewerM 10:00 Welcome. Can you tell us about your background?",
            "",
            "Laura 10:01 Sure! I have been working in HR for about 15 years.",
            "I started in recruitment and moved into organizational development.",
            "Later, I focused on diversity and inclusion initiatives.",
            "",
            "InterviewerM 10:03 That's fascinating. How has AI changed your work?",
            "",
            "Laura 10:04 It has been transformative, honestly.",
            "We now use AI tools for screening resumes, which saves us hours.",
            "",
            "Dana 10:05 I'd like to add to that.",
            "From a governance perspective, we need to be very careful.",
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["InterviewerM", "Laura", "Dana"])
        rows = self._read_csv_rows()

        # Count rows per speaker
        interviewer_rows = [r for r in rows if r['name'] == 'InterviewerM']
        laura_rows = [r for r in rows if r['name'] == 'Laura']
        dana_rows = [r for r in rows if r['name'] == 'Dana']

        self.assertEqual(len(interviewer_rows), 2)
        self.assertEqual(len(laura_rows), 5)  # 3 + 2 continuation
        self.assertEqual(len(dana_rows), 2)  # 1 + 1 continuation

    def test_multiple_interviews_combined(self):
        """Two separate interview files combined into one CSV, each with proper source_file."""
        path1 = self._create_docx_soft_breaks("interview_batch1.docx", [
            "InterviewerM 09:00 Welcome to interview one.",
            "",
            "Laura 09:01 Thank you for having me.",
        ])
        path2 = self._create_docx_soft_breaks("interview_batch2.docx", [
            "InterviewerB 14:00 Let's begin interview two.",
            "",
            "Dana 14:01 Happy to be here.",
        ])
        process_docx_files([path1, path2], self.output_csv, log_callback=self._log,
                           speaker_list=["InterviewerM", "InterviewerB", "Laura", "Dana"])
        rows = self._read_csv_rows()

        batch1_rows = [r for r in rows if r['source_file'] == 'interview_batch1.docx']
        batch2_rows = [r for r in rows if r['source_file'] == 'interview_batch2.docx']
        self.assertEqual(len(batch1_rows), 2)
        self.assertEqual(len(batch2_rows), 2)

    def test_long_statement_preserved_fully(self):
        """Very long statements should not be truncated."""
        long_text = "This is a very detailed response. " * 50  # ~1700 chars
        path = self._create_docx("test.docx", [
            f"Alice 10:00 {long_text.strip()}"
        ])
        process_docx_files([path], self.output_csv, log_callback=self._log,
                           speaker_list=["Alice"])
        rows = self._read_csv_rows()
        self.assertEqual(len(rows), 1)
        self.assertIn("very detailed response", rows[0]['statement'])
        self.assertTrue(len(rows[0]['statement']) > 1000)


if __name__ == '__main__':
    unittest.main()
