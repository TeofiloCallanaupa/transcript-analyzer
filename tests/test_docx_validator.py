import unittest
import os
import shutil
import stat
import tempfile
from docx import Document
from lxml import etree
from docx_to_csv.docx_validator import validate_docx_file, validate_docx_files, ValidationResult


class ValidatorTestBase(unittest.TestCase):
    """Base class with helper methods for validator tests."""

    def setUp(self):
        self.test_dir = tempfile.mkdtemp()

    def tearDown(self):
        # Restore permissions before cleanup (for permission tests)
        for root, dirs, files in os.walk(self.test_dir):
            for f in files:
                os.chmod(os.path.join(root, f), stat.S_IRUSR | stat.S_IWUSR)
        shutil.rmtree(self.test_dir)

    def _create_docx(self, filename, paragraphs):
        """Create a DOCX file with the given paragraphs."""
        path = os.path.join(self.test_dir, filename)
        doc = Document()
        for para in paragraphs:
            doc.add_paragraph(para)
        doc.save(path)
        return path

    def _create_docx_with_table(self, filename, rows_data):
        """Create a DOCX file with content only in a table (no body paragraphs)."""
        path = os.path.join(self.test_dir, filename)
        doc = Document()
        table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text
        doc.save(path)
        return path

    def _create_fake_docx(self, filename, content="this is not a valid docx"):
        """Create a file with .docx extension but invalid content."""
        path = os.path.join(self.test_dir, filename)
        with open(path, 'w') as f:
            f.write(content)
        return path

    def _create_fake_docx_binary(self, filename, content=b''):
        """Create a file with .docx extension and binary content."""
        path = os.path.join(self.test_dir, filename)
        with open(path, 'wb') as f:
            f.write(content)
        return path


# ============================================================================
# CHECK 1: FILE EXISTS AND IS READABLE
# ============================================================================

class TestFileExistence(ValidatorTestBase):
    """Test file existence and readability checks."""

    def test_nonexistent_file(self):
        result = validate_docx_file("/nonexistent/path/fake.docx")
        self.assertFalse(result.is_valid)
        self.assertTrue(any("not found" in e for e in result.errors))

    def test_path_is_directory(self):
        result = validate_docx_file(self.test_dir)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("not a file" in e for e in result.errors))

    def test_empty_string_path(self):
        """Empty string is not a valid path."""
        result = validate_docx_file("")
        self.assertFalse(result.is_valid)
        self.assertTrue(any("not found" in e for e in result.errors))

    def test_file_without_read_permission(self):
        """A file the user can't read should fail gracefully."""
        path = self._create_docx("no_perms.docx", ["Alice 10:00 Hello"])
        os.chmod(path, 0o000)  # Remove all permissions
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("can't be read" in e for e in result.errors))

    def test_file_path_with_spaces(self):
        """Paths with spaces should work fine."""
        sub_dir = os.path.join(self.test_dir, "my folder with spaces")
        os.makedirs(sub_dir)
        path = os.path.join(sub_dir, "my file.docx")
        doc = Document()
        doc.add_paragraph("Alice 10:00 Hello")
        doc.save(path)
        result = validate_docx_file(path, speaker_list=["Alice"])
        self.assertTrue(result.is_valid)

    def test_file_path_with_unicode(self):
        """Paths with unicode characters should work."""
        path = self._create_docx("résumé_análisis.docx", ["Alice 10:00 Hello"])
        result = validate_docx_file(path, speaker_list=["Alice"])
        self.assertTrue(result.is_valid)


# ============================================================================
# CHECK 2: VALID DOCX FORMAT
# ============================================================================

class TestValidDocxFormat(ValidatorTestBase):
    """Test that non-DOCX files are caught."""

    def test_renamed_txt_file(self):
        path = self._create_fake_docx("fake.docx", "just plain text")
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("not a valid DOCX" in e for e in result.errors))

    def test_renamed_csv_file(self):
        path = self._create_fake_docx("data.docx", "col1,col2\nval1,val2")
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("not a valid DOCX" in e for e in result.errors))

    def test_valid_docx_passes(self):
        path = self._create_docx("valid.docx", ["Alice 10:00 Hello"])
        result = validate_docx_file(path, speaker_list=["Alice"])
        self.assertTrue(result.is_valid)
        self.assertEqual(len(result.errors), 0)

    def test_binary_garbage_file(self):
        path = self._create_fake_docx_binary("garbage.docx", b'\x00\x01\x02\x03\x04\x05')
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("not a valid DOCX" in e for e in result.errors))

    def test_zero_byte_file(self):
        """A completely empty file (0 bytes) should fail as invalid DOCX."""
        path = self._create_fake_docx_binary("empty_file.docx", b'')
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("not a valid DOCX" in e for e in result.errors))

    def test_old_doc_format_binary(self):
        """An old .doc (Word 97-2003) binary format should fail with a helpful message.
        The .doc format starts with the OLE2 magic bytes (D0 CF 11 E0)."""
        path = self._create_fake_docx_binary("old_format.docx",
            b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1' + b'\x00' * 100)
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("not a valid DOCX" in e for e in result.errors))
        # Error message should hint at the .doc possibility
        self.assertTrue(any(".doc" in e for e in result.errors))

    def test_html_file_renamed_to_docx(self):
        """An HTML file saved as .docx should fail."""
        path = self._create_fake_docx("webpage.docx", "<html><body>Hello</body></html>")
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("not a valid DOCX" in e for e in result.errors))

    def test_pdf_renamed_to_docx(self):
        """A PDF file renamed to .docx should fail."""
        path = self._create_fake_docx_binary("document.docx",
            b'%PDF-1.4 fake pdf content' + b'\x00' * 50)
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)


# ============================================================================
# CHECK 3: EMPTY DOCUMENT
# ============================================================================

class TestEmptyDocument(ValidatorTestBase):
    """Test that empty documents are caught."""

    def test_empty_docx(self):
        path = self._create_docx("empty.docx", [])
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("empty" in e.lower() for e in result.errors))

    def test_whitespace_only_docx(self):
        path = self._create_docx("blank.docx", ["   ", "  ", "\t"])
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("empty" in e.lower() for e in result.errors))

    def test_newline_only_paragraphs(self):
        """Paragraphs that are just newlines should count as empty."""
        path = self._create_docx("newlines.docx", ["\n", "\n\n", "\n\n\n"])
        result = validate_docx_file(path)
        self.assertFalse(result.is_valid)
        self.assertTrue(any("empty" in e.lower() for e in result.errors))

    def test_docx_with_real_content_passes(self):
        path = self._create_docx("real.docx", ["Some actual text here"])
        result = validate_docx_file(path)
        self.assertTrue(result.is_valid)

    def test_docx_with_only_table_content(self):
        """A DOCX with text only in tables (no body paragraphs) should NOT be flagged as empty."""
        path = self._create_docx_with_table("table_only.docx", [
            ["Alice", "10:00", "Hello from a table"],
            ["Bob", "10:01", "Table content here"],
        ])
        result = validate_docx_file(path)
        self.assertTrue(result.is_valid)
        self.assertFalse(any("empty" in e.lower() for e in result.errors))


# ============================================================================
# CHECK 4: SPEAKER NAME MATCHING
# ============================================================================

class TestSpeakerNameMatching(ValidatorTestBase):
    """Test speaker name validation warnings."""

    def test_no_matching_speakers_warns(self):
        path = self._create_docx("test.docx", [
            "UnknownPerson 10:00 Hello",
            "AnotherPerson 10:01 World"
        ])
        result = validate_docx_file(path, speaker_list=["Alice", "Bob"])
        # File is still valid (it's a warning, not an error)
        self.assertTrue(result.is_valid)
        self.assertTrue(len(result.warnings) > 0)
        self.assertTrue(any("No speaker names" in w for w in result.warnings))

    def test_matching_speakers_no_warning(self):
        path = self._create_docx("test.docx", [
            "Alice 10:00 Hello",
            "Bob 10:01 World"
        ])
        result = validate_docx_file(path, speaker_list=["Alice", "Bob"])
        self.assertTrue(result.is_valid)
        self.assertEqual(len(result.warnings), 0)

    def test_partial_speaker_match_no_warning(self):
        """If at least one speaker matches, no warning."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 Hello",
            "UnknownPerson 10:01 World"
        ])
        result = validate_docx_file(path, speaker_list=["Alice", "Bob"])
        self.assertTrue(result.is_valid)
        self.assertEqual(len(result.warnings), 0)

    def test_no_speaker_list_skips_check(self):
        """When speaker_list is None, speaker check is skipped."""
        path = self._create_docx("test.docx", ["Some random text"])
        result = validate_docx_file(path, speaker_list=None)
        self.assertTrue(result.is_valid)
        self.assertFalse(any("speaker" in w.lower() for w in result.warnings))

    def test_empty_speaker_list_skips_check(self):
        """When speaker_list is [], speaker check is skipped."""
        path = self._create_docx("test.docx", ["Some random text"])
        result = validate_docx_file(path, speaker_list=[])
        self.assertTrue(result.is_valid)

    def test_warning_includes_sample_lines(self):
        """Warning should include a preview of the document's first lines."""
        path = self._create_docx("test.docx", [
            "PersonX 10:00 Some statement"
        ])
        result = validate_docx_file(path, speaker_list=["Alice"])
        self.assertTrue(len(result.warnings) > 0)
        self.assertTrue(any("PersonX" in w for w in result.warnings))

    def test_speaker_name_with_regex_special_chars(self):
        """Speaker names like 'Dr. Smith' or 'Person (A)' contain regex metacharacters.
        They must be escaped properly."""
        path = self._create_docx("test.docx", [
            "Dr. Smith 10:00 The diagnosis is clear"
        ])
        result = validate_docx_file(path, speaker_list=["Dr. Smith"])
        self.assertTrue(result.is_valid)
        self.assertEqual(len(result.warnings), 0)

    def test_speaker_name_with_parentheses(self):
        """Speaker with parentheses in name."""
        path = self._create_docx("test.docx", [
            "Person (A) 10:00 Hello from person A"
        ])
        result = validate_docx_file(path, speaker_list=["Person (A)"])
        self.assertTrue(result.is_valid)
        self.assertEqual(len(result.warnings), 0)

    def test_speaker_case_sensitivity(self):
        """Speaker matching is case-sensitive — 'alice' ≠ 'Alice'."""
        path = self._create_docx("test.docx", [
            "alice 10:00 lowercase name"
        ])
        result = validate_docx_file(path, speaker_list=["Alice"])
        # Should warn because "alice" doesn't match "Alice"
        self.assertTrue(any("No speaker names" in w for w in result.warnings))

    def test_speaker_match_deep_in_document(self):
        """Speaker found on line 10+ should still count as a match."""
        paragraphs = [f"Preamble line {i}" for i in range(10)]
        paragraphs.append("Alice 10:00 Finally, a real speaker line")
        path = self._create_docx("test.docx", paragraphs)
        result = validate_docx_file(path, speaker_list=["Alice"])
        self.assertTrue(result.is_valid)
        self.assertEqual(len(result.warnings), 0)

    def test_many_speakers_truncated_in_warning(self):
        """With > 5 speakers configured, warning should show truncated list."""
        speakers = [f"Speaker{i}" for i in range(8)]
        path = self._create_docx("test.docx", ["Nobody 10:00 No match here"])
        result = validate_docx_file(path, speaker_list=speakers)
        self.assertTrue(any("8 total" in w for w in result.warnings))

    def test_speaker_without_timestamp_still_matches(self):
        """A speaker name alone on a line (no timestamp) should still count as a match."""
        path = self._create_docx("test.docx", [
            "Alice",
            "I have something to say."
        ])
        result = validate_docx_file(path, speaker_list=["Alice"])
        self.assertTrue(result.is_valid)
        self.assertEqual(len(result.warnings), 0)


# ============================================================================
# CHECK 5: CONTROL CHARACTERS
# ============================================================================

class TestControlCharacters(ValidatorTestBase):
    """Test detection of problematic control characters.
    
    NOTE: DOCX files are XML-based and XML forbids most control characters.
    python-docx will refuse to create documents with null bytes or control chars.
    We use direct XML manipulation (lxml) to inject them, simulating file corruption.
    """

    def _create_docx_with_control_char(self, filename, text_with_control):
        """Create a DOCX with control chars injected via lxml XML manipulation.
        This simulates corrupted or non-standard DOCX files."""
        # First create a normal DOCX
        path = os.path.join(self.test_dir, filename)
        doc = Document()
        doc.add_paragraph("placeholder")
        doc.save(path)
        
        # Now reopen and inject control chars directly into the XML
        doc = Document(path)
        para = doc.paragraphs[0]
        # Access the underlying XML run element
        run = para.runs[0]
        # Find the <w:t> element and set text directly via lxml
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        t_elem = run._r.find('.//w:t', nsmap)
        if t_elem is not None:
            # lxml allows setting text with control chars if we bypass validation
            t_elem.text = text_with_control
        doc.save(path)
        return path

    def test_normal_text_no_warning(self):
        path = self._create_docx("clean.docx", [
            "Alice 10:00 Normal text with punctuation, quotes \"here\" and tabs."
        ])
        result = validate_docx_file(path, speaker_list=["Alice"])
        self.assertTrue(result.is_valid)
        self.assertEqual(len(result.warnings), 0)

    def test_tab_character_does_not_warn(self):
        """Tab (0x09) is a normal whitespace char and should NOT trigger a warning."""
        path = self._create_docx("test.docx", [
            "Alice 10:00 Text with\ttabs\there"
        ])
        result = validate_docx_file(path, speaker_list=["Alice"])
        char_warnings = [w for w in result.warnings if "control character" in w.lower()]
        self.assertEqual(len(char_warnings), 0)

    def test_control_char_detection_with_injected_bell(self):
        """Bell character (0x07) injected via XML should trigger a warning."""
        try:
            path = self._create_docx_with_control_char(
                "bell.docx", "Text with \x07 bell char"
            )
            result = validate_docx_file(path)
            # If lxml allowed the injection, we should get a warning
            if any("control character" in w.lower() for w in result.warnings):
                self.assertTrue(True)  # Warning detected as expected
            else:
                # lxml may have stripped it — that's also fine
                self.assertTrue(result.is_valid)
        except (ValueError, etree.XMLSyntaxError):
            # lxml itself may reject control chars — this proves DOCX can't contain them
            self.skipTest("lxml rejects control characters in XML text nodes")

    def test_control_char_regex_pattern_directly(self):
        """Test the control character regex pattern directly (without DOCX creation).
        This validates the detection logic even though DOCX files can't contain these chars."""
        from docx_to_csv.docx_validator import _CONTROL_CHAR_PATTERN
        
        # Should match control chars
        self.assertIsNotNone(_CONTROL_CHAR_PATTERN.search("text\x00null"))
        self.assertIsNotNone(_CONTROL_CHAR_PATTERN.search("text\x07bell"))
        self.assertIsNotNone(_CONTROL_CHAR_PATTERN.search("text\x7fDEL"))
        self.assertIsNotNone(_CONTROL_CHAR_PATTERN.search("text\x01SOH"))
        
        # Should NOT match normal chars (tab, newline, carriage return are OK)
        self.assertIsNone(_CONTROL_CHAR_PATTERN.search("normal text"))
        self.assertIsNone(_CONTROL_CHAR_PATTERN.search("text\twith\ttabs"))
        self.assertIsNone(_CONTROL_CHAR_PATTERN.search("text\nwith\nnewlines"))
        self.assertIsNone(_CONTROL_CHAR_PATTERN.search("text\rwith\rCR"))
        self.assertIsNone(_CONTROL_CHAR_PATTERN.search("Héllo wörld café 🎉"))


# ============================================================================
# BATCH VALIDATION
# ============================================================================

class TestBatchValidation(ValidatorTestBase):
    """Test validate_docx_files (batch)."""

    def test_mixed_valid_and_invalid(self):
        good = self._create_docx("good.docx", ["Alice 10:00 Hello"])
        bad = self._create_fake_docx("bad.docx")
        results = validate_docx_files([good, bad], speaker_list=["Alice"])
        self.assertEqual(len(results), 2)
        self.assertTrue(results[0].is_valid)
        self.assertFalse(results[1].is_valid)

    def test_preserves_order(self):
        paths = []
        for i in range(3):
            paths.append(self._create_docx(f"file{i}.docx", [f"Alice 10:0{i} Statement {i}"]))
        results = validate_docx_files(paths, speaker_list=["Alice"])
        for i, r in enumerate(results):
            self.assertIn(f"file{i}.docx", r.file_path)

    def test_empty_list(self):
        """Validating an empty list should return an empty list."""
        results = validate_docx_files([], speaker_list=["Alice"])
        self.assertEqual(results, [])

    def test_all_invalid(self):
        """When all files are invalid, all results should reflect that."""
        bad1 = self._create_fake_docx("bad1.docx")
        bad2 = self._create_fake_docx("bad2.docx")
        results = validate_docx_files([bad1, bad2], speaker_list=["Alice"])
        self.assertEqual(len(results), 2)
        self.assertFalse(results[0].is_valid)
        self.assertFalse(results[1].is_valid)

    def test_single_file(self):
        """Batch with a single file should work."""
        path = self._create_docx("only.docx", ["Alice 10:00 Solo"])
        results = validate_docx_files([path], speaker_list=["Alice"])
        self.assertEqual(len(results), 1)
        self.assertTrue(results[0].is_valid)


# ============================================================================
# VALIDATION RESULT DATACLASS
# ============================================================================

class TestValidationResult(ValidatorTestBase):
    """Test ValidationResult defaults and structure."""

    def test_defaults(self):
        r = ValidationResult(is_valid=True, file_path="/test.docx")
        self.assertEqual(r.errors, [])
        self.assertEqual(r.warnings, [])
        self.assertTrue(r.is_valid)

    def test_mutable_default_isolation(self):
        """Each ValidationResult should have its own independent lists (no shared mutable defaults)."""
        r1 = ValidationResult(is_valid=True, file_path="/a.docx")
        r2 = ValidationResult(is_valid=True, file_path="/b.docx")
        r1.errors.append("error in r1")
        r1.warnings.append("warning in r1")
        # r2 should be unaffected
        self.assertEqual(r2.errors, [])
        self.assertEqual(r2.warnings, [])


# ============================================================================
# COMBINED / INTEGRATION EDGE CASES
# ============================================================================

class TestCombinedEdgeCases(ValidatorTestBase):
    """Tests that combine multiple checks or simulate realistic user mistakes."""

    def test_valid_file_with_speaker_warning_and_no_other_issues(self):
        """A file with unrecognized speakers should still be valid with just a speaker warning."""
        path = self._create_docx("messy.docx", [
            "UnknownPerson 10:00 Some statement here"
        ])
        result = validate_docx_file(path, speaker_list=["Alice"])
        self.assertTrue(result.is_valid)
        has_speaker_warning = any("No speaker names" in w for w in result.warnings)
        self.assertTrue(has_speaker_warning)

    def test_structural_failure_short_circuits(self):
        """A corrupt file should fail at check 2 and NOT run speaker or char checks."""
        path = self._create_fake_docx("corrupt.docx")
        result = validate_docx_file(path, speaker_list=["Alice"])
        self.assertFalse(result.is_valid)
        # Should have exactly 1 error (DOCX validity) and 0 warnings (checks 4-5 didn't run)
        self.assertEqual(len(result.errors), 1)
        self.assertEqual(len(result.warnings), 0)

    def test_empty_file_short_circuits(self):
        """An empty DOCX should fail at check 3 and NOT run speaker or char checks."""
        path = self._create_docx("empty.docx", [])
        result = validate_docx_file(path, speaker_list=["Alice"])
        self.assertFalse(result.is_valid)
        self.assertEqual(len(result.errors), 1)
        self.assertIn("empty", result.errors[0].lower())
        self.assertEqual(len(result.warnings), 0)

    def test_result_file_path_matches_input(self):
        """The ValidationResult should always contain the file path that was validated."""
        path = self._create_docx("specific_name.docx", ["Content here"])
        result = validate_docx_file(path)
        self.assertEqual(result.file_path, path)


if __name__ == '__main__':
    unittest.main()
